import docx
d = docx.Document('CL_KC_N26 copy.docx')
d.paragraphs
# print(d.paragraphs[0].text)
# print(d.paragraphs[1].text)
p = d.paragraphs[1]
p2 = d.paragraphs[4]
print(p2)
# print(p.runs)
# print(p.runs[0].text)
# print(p.runs[1].text)
# # print(p.runs[2].text)
# print(p.runs[0].bold)
# print(p.runs[0].bold == None)
# print(p.runs[0].italic)
# print(p.runs[0].underline)

p.runs[1].text = 'italic and underline'
d.save('demo2.docx')
# print(p.style)
# p2.style = 'Strong'  # has some proble here: 
# ValueError: assigned style is type CHARACTER (2), need type PARAGRAPH (1)
# # p2.style = d.styles['Strong']
# d.save('demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('This is another paragraph')
d.save('demo4.docx')
p = d.paragraphs[0]
p.add_run('This is a new run.')
# p.runs
p.runs[1].bold = True
d.save('demo5.docx')

d = docx.Document()




